"""
文档读取器 - 支持Excel/Word/PDF/OCR多种格式
将各种格式文档转换为统一的文本块序列
"""

import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import re

from .schemas import DocumentInput, TextBlock, SourceChannel


logger = logging.getLogger(__name__)


class DocumentReader:
    """文档读取器"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.ocr_enabled = self.config.get("ocr", {}).get("enabled", False)
        
        # 延迟导入可选依赖
        self._pandas = None
        self._docx = None
        self._pdfplumber = None
        self._paddleocr = None
    
    def read_document(self, file_path: str, use_ocr: bool = False) -> DocumentInput:
        """
        读取文档并转换为统一格式
        
        Args:
            file_path: 文件路径
            use_ocr: 是否使用OCR
            
        Returns:
            文档输入对象
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        try:
            if suffix == '.xlsx':
                return self._read_excel(file_path)
            elif suffix == '.docx':
                return self._read_word(file_path)
            elif suffix == '.pdf':
                if use_ocr:
                    return self._read_pdf_ocr(file_path)
                else:
                    return self._read_pdf_text(file_path)
            elif suffix in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
                return self._read_image_ocr(file_path)
            else:
                # 尝试作为文本文件读取
                return self._read_text(file_path)
                
        except Exception as e:
            logger.error(f"文档读取失败 {file_path}: {e}")
            raise
    
    def _read_excel(self, file_path: str) -> DocumentInput:
        """读取Excel文件"""
        if self._pandas is None:
            import pandas as pd
            self._pandas = pd
        
        blocks = []
        line_no = 1
        
        try:
            # 读取所有工作表
            excel_file = self._pandas.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = self._pandas.read_excel(file_path, sheet_name=sheet_name)
                
                # 添加工作表标题
                blocks.append(TextBlock(
                    type="sheet_title",
                    text=f"工作表: {sheet_name}",
                    line_no=line_no
                ))
                line_no += 1
                
                # 处理每一行
                for idx, row in df.iterrows():
                    # 合并非空单元格
                    row_text = " ".join(str(cell) for cell in row.values if pd.notna(cell))
                    
                    if row_text.strip():
                        blocks.append(TextBlock(
                            type="text",
                            text=row_text.strip(),
                            line_no=line_no
                        ))
                        line_no += 1
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.EXCEL,
                    "sheets": excel_file.sheet_names,
                    "total_lines": line_no - 1
                }
            )
            
        except Exception as e:
            logger.error(f"Excel读取失败: {e}")
            raise
    
    def _read_word(self, file_path: str) -> DocumentInput:
        """读取Word文件"""
        if self._docx is None:
            from docx import Document
            self._docx = Document
        
        blocks = []
        line_no = 1
        
        try:
            doc = self._docx(file_path)
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    blocks.append(TextBlock(
                        type="text",
                        text=text,
                        line_no=line_no
                    ))
                    line_no += 1
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        blocks.append(TextBlock(
                            type="table_row",
                            text=row_text,
                            line_no=line_no
                        ))
                        line_no += 1
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.WORD,
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "total_lines": line_no - 1
                }
            )
            
        except Exception as e:
            logger.error(f"Word读取失败: {e}")
            raise
    
    def _read_pdf_text(self, file_path: str) -> DocumentInput:
        """读取PDF文件（文本模式）"""
        if self._pdfplumber is None:
            import pdfplumber
            self._pdfplumber = pdfplumber
        
        blocks = []
        line_no = 1
        
        try:
            with self._pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 添加页面标题
                    blocks.append(TextBlock(
                        type="page_title",
                        text=f"第{page_num}页",
                        line_no=line_no
                    ))
                    line_no += 1
                    
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        # 按行分割
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                blocks.append(TextBlock(
                                    type="text",
                                    text=line,
                                    line_no=line_no
                                ))
                                line_no += 1
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.PDF,
                    "pages": len(pdf.pages),
                    "total_lines": line_no - 1
                }
            )
            
        except Exception as e:
            logger.error(f"PDF读取失败: {e}")
            raise
    
    def _read_pdf_ocr(self, file_path: str) -> DocumentInput:
        """读取PDF文件（OCR模式）"""
        if self._paddleocr is None:
            from paddleocr import PaddleOCR
            self._paddleocr = PaddleOCR(use_angle_cls=True, lang='ch')
        
        blocks = []
        line_no = 1
        
        try:
            # 将PDF转换为图片然后OCR
            # 这里简化处理，实际需要pdf2image等工具
            
            # 占位符实现
            blocks.append(TextBlock(
                type="text",
                text="OCR功能待实现",
                line_no=line_no,
                ocr_conf=0.9
            ))
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.OCR,
                    "ocr_engine": "paddleocr",
                    "total_lines": len(blocks)
                }
            )
            
        except Exception as e:
            logger.error(f"PDF OCR失败: {e}")
            raise
    
    def _read_image_ocr(self, file_path: str) -> DocumentInput:
        """读取图片文件（OCR）"""
        if self._paddleocr is None:
            from paddleocr import PaddleOCR
            self._paddleocr = PaddleOCR(use_angle_cls=True, lang='ch')
        
        blocks = []
        
        try:
            result = self._paddleocr.ocr(file_path, cls=True)
            
            for line_result in result[0]:  # result是三层嵌套列表
                bbox, (text, confidence) = line_result
                
                if confidence > 0.5:  # 置信度阈值
                    blocks.append(TextBlock(
                        type="text",
                        text=text,
                        line_no=len(blocks) + 1,
                        ocr_conf=confidence,
                        bbox=bbox
                    ))
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.OCR,
                    "ocr_engine": "paddleocr",
                    "total_lines": len(blocks)
                }
            )
            
        except Exception as e:
            logger.error(f"图片OCR失败: {e}")
            raise
    
    def _read_text(self, file_path: str) -> DocumentInput:
        """读取纯文本文件"""
        blocks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                for line_no, line in enumerate(f, 1):
                    line = line.strip()
                    if line:
                        blocks.append(TextBlock(
                            type="text",
                            text=line,
                            line_no=line_no
                        ))
            
            return DocumentInput(
                source_id=f"file://{file_path}",
                blocks=blocks,
                meta={
                    "channel": SourceChannel.TEXT,
                    "total_lines": len(blocks)
                }
            )
            
        except Exception as e:
            logger.error(f"文本读取失败: {e}")
            raise
    
    def batch_read_documents(self, file_paths: List[str], use_ocr: bool = False) -> List[DocumentInput]:
        """
        批量读取文档
        
        Args:
            file_paths: 文件路径列表
            use_ocr: 是否使用OCR
            
        Returns:
            文档列表
        """
        documents = []
        
        for file_path in file_paths:
            try:
                document = self.read_document(file_path, use_ocr)
                documents.append(document)
                logger.info(f"文档读取成功: {file_path}")
            except Exception as e:
                logger.error(f"文档读取失败 {file_path}: {e}")
                continue
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        formats = ['.xlsx', '.docx', '.pdf', '.txt']
        
        if self.ocr_enabled:
            formats.extend(['.png', '.jpg', '.jpeg', '.bmp', '.tiff'])
        
        return formats


def read_document(file_path: str, use_ocr: bool = False, config: Dict[str, Any] = None) -> DocumentInput:
    """
    便捷函数：读取单个文档
    
    Args:
        file_path: 文件路径
        use_ocr: 是否使用OCR
        config: 配置参数
        
    Returns:
        文档输入对象
    """
    reader = DocumentReader(config)
    return reader.read_document(file_path, use_ocr)


# Word解析优化代码
# 
    def _merge_text_blocks_smart(self, blocks):
        """智能合并文本块"""
        if not blocks:
            return blocks
        
        merged = []
        current_block = None
        
        for block in blocks:
            text = block.text.strip()
            if not text:
                continue
            
            # 判断是否应该合并
            should_merge = False
            
            if current_block:
                current_text = current_block.text
                
                # 合并条件
                if (len(current_text) < 50 and len(text) < 100 and 
                    not any(text.startswith(prefix) for prefix in ['A.', 'B.', 'C.', 'D.', 'E.', 'F.']) and
                    not '答案' in text and not '解析' in text):
                    should_merge = True
            
            if should_merge:
                # 合并到当前块
                current_block.text = current_block.text + " " + text
            else:
                # 保存当前块，开始新块
                if current_block:
                    merged.append(current_block)
                current_block = block
        
        # 添加最后一个块
        if current_block:
            merged.append(current_block)
        
        return merged
    
