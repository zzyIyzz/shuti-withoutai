#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word题库解析模块
支持多种Word题库格式
"""

import re
from pathlib import Path
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None

class WordTikuParser:
    def __init__(self, word_file):
        self.word_file = Path(word_file)
        self.questions = []
    
    def parse(self):
        """解析Word题库"""
        if Document is None:
            print("错误：需要安装 python-docx 库")
            print("运行: pip install python-docx")
            return []
        
        try:
            doc = Document(self.word_file)
            
            # 尝试多种解析方式
            # 方式1：表格格式
            questions_from_tables = self.parse_tables(doc)
            if questions_from_tables:
                print(f"从表格中解析到 {len(questions_from_tables)} 道题")
                return questions_from_tables
            
            # 方式2：段落格式（结构化）
            questions_from_paragraphs = self.parse_structured_paragraphs(doc)
            if questions_from_paragraphs:
                print(f"从结构化段落中解析到 {len(questions_from_paragraphs)} 道题")
                return questions_from_paragraphs
            
            # 方式3：段落格式（自由文本）
            questions_from_text = self.parse_free_text(doc)
            if questions_from_text:
                print(f"从文本中解析到 {len(questions_from_text)} 道题")
                return questions_from_text
            
            print("警告：未能识别Word题库格式")
            return []
            
        except Exception as e:
            print(f"解析Word文件失败: {e}")
            return []
    
    def parse_tables(self, doc):
        """解析表格格式的题库"""
        questions = []
        
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            
            # 读取表头
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
            
            # 识别列
            col_map = self.identify_columns(headers)
            
            if not col_map.get('question') or not col_map.get('answer'):
                continue
            
            # 读取题目行
            for row_idx, row in enumerate(table.rows[1:], start=1):
                cells = [cell.text.strip() for cell in row.cells]
                
                if col_map['question'] >= len(cells):
                    continue
                
                question_text = cells[col_map['question']]
                if not question_text:
                    continue
                
                question = {
                    'id': len(questions) + 1,
                    'question': question_text,
                    'answer': '',
                    'options': {},
                    'type': 'unknown',
                    'explanation': ''
                }
                
                # 答案
                if col_map.get('answer') and col_map['answer'] < len(cells):
                    question['answer'] = cells[col_map['answer']]
                
                # 选项
                for opt_key, opt_idx in col_map.get('options', {}).items():
                    if opt_idx < len(cells) and cells[opt_idx]:
                        question['options'][opt_key] = cells[opt_idx]
                
                # 题型
                if col_map.get('type') and col_map['type'] < len(cells):
                    question['type'] = cells[col_map['type']]
                else:
                    question['type'] = self.detect_question_type(question)
                
                # 解析
                if col_map.get('explanation') and col_map['explanation'] < len(cells):
                    question['explanation'] = cells[col_map['explanation']]
                
                questions.append(question)
        
        return questions
    
    def parse_structured_paragraphs(self, doc):
        """
        解析结构化段落格式
        例如：
        1. 题目内容
        A. 选项A
        B. 选项B
        答案：A
        """
        questions = []
        current_question = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测题目开始（数字编号）
            question_match = re.match(r'^(\d+)[、.．]\s*(.+)', text)
            if question_match:
                # 保存上一题
                if current_question and current_question.get('question'):
                    questions.append(current_question)
                
                # 开始新题
                current_question = {
                    'id': int(question_match.group(1)),
                    'question': question_match.group(2),
                    'answer': '',
                    'options': {},
                    'type': 'unknown',
                    'explanation': ''
                }
                continue
            
            if not current_question:
                continue
            
            # 检测选项
            option_match = re.match(r'^([A-F])[、.．)\s]\s*(.+)', text)
            if option_match:
                opt_key = option_match.group(1)
                opt_value = option_match.group(2)
                current_question['options'][opt_key] = opt_value
                continue
            
            # 检测答案
            answer_match = re.match(r'^(?:答案|正确答案)[：:]\s*(.+)', text, re.IGNORECASE)
            if answer_match:
                current_question['answer'] = answer_match.group(1).strip()
                continue
            
            # 检测题型
            type_match = re.match(r'^(?:题型|类型)[：:]\s*(.+)', text)
            if type_match:
                current_question['type'] = type_match.group(1).strip()
                continue
            
            # 检测解析
            exp_match = re.match(r'^(?:解析|解释|说明)[：:]\s*(.+)', text)
            if exp_match:
                current_question['explanation'] = exp_match.group(1).strip()
                continue
        
        # 保存最后一题
        if current_question and current_question.get('question'):
            # 自动判断题型
            if current_question['type'] == 'unknown':
                current_question['type'] = self.detect_question_type(current_question)
            questions.append(current_question)
        
        return questions
    
    def parse_free_text(self, doc):
        """
        解析自由文本格式
        尝试智能识别题目边界
        """
        questions = []
        all_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # 按题号分割
        # 支持格式：1. 1、 一、 （1） [1]
        patterns = [
            r'\n(\d+)[、.．]\s*',
            r'\n(\d+)\)\s*',
            r'\n\((\d+)\)\s*',
            r'\n\[(\d+)\]\s*',
        ]
        
        for pattern in patterns:
            parts = re.split(pattern, all_text)
            if len(parts) > 3:  # 找到了分割点
                for i in range(1, len(parts), 2):
                    if i + 1 < len(parts):
                        q_num = parts[i]
                        q_content = parts[i + 1].strip()
                        
                        question = self.parse_single_question_text(q_content, int(q_num))
                        if question:
                            questions.append(question)
                
                if questions:
                    break
        
        return questions
    
    def parse_single_question_text(self, text, q_id):
        """解析单个题目的文本"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return None
        
        question = {
            'id': q_id,
            'question': '',
            'answer': '',
            'options': {},
            'type': 'unknown',
            'explanation': ''
        }
        
        # 第一行是题目
        question['question'] = lines[0]
        
        # 解析其他行
        for line in lines[1:]:
            # 选项
            opt_match = re.match(r'^([A-F])[、.．)\s]\s*(.+)', line)
            if opt_match:
                question['options'][opt_match.group(1)] = opt_match.group(2)
                continue
            
            # 答案
            ans_match = re.match(r'^(?:答案|正确答案)[：:]\s*(.+)', line, re.IGNORECASE)
            if ans_match:
                question['answer'] = ans_match.group(1).strip()
                continue
            
            # 解析
            exp_match = re.match(r'^(?:解析|解释|说明)[：:]\s*(.+)', line)
            if exp_match:
                question['explanation'] = exp_match.group(1).strip()
                continue
        
        # 自动判断题型
        question['type'] = self.detect_question_type(question)
        
        return question if question['answer'] else None
    
    def identify_columns(self, headers):
        """智能识别列（与Excel解析相同）"""
        col_map = {}
        
        # 题目列
        question_keywords = ['题目', '试题', '问题', '题干', 'question', 'content']
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if any(kw in h_lower for kw in question_keywords):
                col_map['question'] = i
                break
        
        # 答案列
        answer_keywords = ['答案', '正确答案', 'answer', '标准答案']
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if any(kw in h_lower for kw in answer_keywords):
                col_map['answer'] = i
                break
        
        # 选项列
        option_map = {}
        for opt in ['A', 'B', 'C', 'D', 'E', 'F']:
            for i, h in enumerate(headers):
                if h.strip() == opt or h.strip() == f'选项{opt}':
                    option_map[opt] = i
                    break
        if option_map:
            col_map['options'] = option_map
        
        # 题型列
        type_keywords = ['题型', '类型', 'type']
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if any(kw in h_lower for kw in type_keywords):
                col_map['type'] = i
                break
        
        # 解析列
        exp_keywords = ['解析', '解释', '说明', 'explanation']
        for i, h in enumerate(headers):
            h_lower = h.lower()
            if any(kw in h_lower for kw in exp_keywords):
                col_map['explanation'] = i
                break
        
        return col_map
    
    def detect_question_type(self, question):
        """题型识别 - 使用智能识别引擎"""
        try:
            from 智能题型识别 import detect_question_type
            return detect_question_type(
                question['question'],
                question['answer'],
                question.get('options')
            )
        except:
            return self._simple_detect(question)
    
    def _simple_detect(self, question):
        """备用题型识别方法 - 详细完备版（与题库管理模块保持一致）"""
        import re
        
        q_text = question['question']
        answer = str(question['answer']).strip().upper()
        q_lower = q_text.lower()
        
        # ========== 第一优先级：判断题识别 ==========
        判断题答案集 = {'对', '错', '√', '×', 'T', 'F', 'TRUE', 'FALSE', '正确', '错误', '是', '否'}
        if answer in 判断题答案集:
            return '判断题'
        if '对错' in answer or ('对' in answer and '错' in answer):
            return '判断题'
        
        判断题关键词 = ['是否正确', '对吗', '对么', '错吗', '错么', '判断题', '判断正误', 
                    '说法正确', '说法错误', '表述正确', '表述错误',
                    '观点正确', '观点错误', '是否准确', '对不对', '错不错']
        if any(keyword in q_lower for keyword in 判断题关键词):
            return '判断题'
        
        # ========== 第二优先级：选择题识别 ==========
        选项模式列表 = [
            r'[A-F][、\.]\s*\S+',
            r'[A-F]\s*[、\.]\s*\S+',
        ]
        
        has_options = False
        for 模式 in 选项模式列表:
            if re.search(模式, q_text):
                has_options = True
                break
        
        if question.get('options'):
            has_options = True
        
        # 多选题识别
        if has_options or '多选' in q_lower or '多项' in q_lower:
            if len(answer) > 1 and all(c in 'ABCDEF' for c in answer):
                return '多选题'
            
            多选关键词 = ['多选', '多项选择', '多项', '哪些', '哪几个', '哪几项', 
                      '包括哪些', '包含哪些', '有哪些', '正确的有', '错误的有',
                      '以下正确', '以下错误', '下列正确', '下列错误']
            if any(keyword in q_lower for keyword in 多选关键词):
                return '多选题'
        
        # 单选题识别
        if has_options:
            if len(answer) == 1 and answer in 'ABCDEF':
                return '单选题'
            
            单选关键词 = ['单选', '单项选择', '单项', '哪个', '哪项', '哪一',
                      '最正确', '最合适', '最恰当', '最准确', '最合理']
            if any(keyword in q_lower for keyword in 单选关键词):
                return '单选题'
            
            return '单选题'
        
        # 无选项但答案是字母
        if len(answer) > 1 and all(c in 'ABCDEF' for c in answer):
            return '多选题'
        if len(answer) == 1 and answer in 'ABCDEF':
            return '单选题'
        
        # ========== 第三优先级：简答题识别 ==========
        if len(answer) > 20 and any(p in answer for p in ['。', '，', '；', '、']):
            return '简答题'
        
        简答题关键词 = ['简述', '简要说明', '说明', '论述', '阐述', '分析', '解释', 
                    '描述', '叙述', '介绍', '回答', '解答', '谈谈', '试述',
                    '如何', '怎样', '怎么', '为什么', '为何',
                    '什么是', '何为', '定义', '概念', '含义']
        if any(keyword in q_lower for keyword in 简答题关键词):
            if len(answer) > 5:
                return '简答题'
        
        # ========== 第四优先级：填空题识别 ==========
        填空标记 = ['_', '____', '（）', '()', '【】', '[]']
        if any(mark in q_text for mark in 填空标记):
            return '填空题'
        
        填空句式 = [
            r'是\s*[（(]?\s*[）)]?',
            r'为\s*[（(]?\s*[）)]?',
            r'应\s*[（(]?\s*[）)]?',
            r'等于\s*[（(]?\s*[）)]?',
        ]
        if any(re.search(pattern, q_text) for pattern in 填空句式):
            return '填空题'
        
        技术单位 = ['MPa', 'KPA', 'PA', 'KV', 'V', 'MV', 'A', 'MA', 'KA',
                  'W', 'KW', 'MW', 'HZ', 'KHZ', 'MHZ', 'Ω', 'KΩ', 'MΩ',
                  '米', '厘米', '毫米', '千米', '克', '千克', '吨', '升', '毫升',
                  '度', '℃', '%', '‰', '年', '月', '日', '天', '小时', '分钟', '秒']
        
        if any(char.isdigit() for char in answer):
            if any(unit in answer for unit in 技术单位):
                return '填空题'
            if re.search(r'\d+', answer):
                return '填空题'
        
        填空关键词 = ['等于', '约为', '大约', '标准', '规定', '要求',
                    '必须', '应该', '需要', '达到', '超过', '低于',
                    '称为', '叫做', '简称']
        if any(keyword in q_lower for keyword in 填空关键词):
            if not (len(answer) <= 2 and all(c in 'ABCDEF' for c in answer)):
                return '填空题'
        
        # ========== 最终兜底策略 ==========
        if len(answer) > 15:
            return '简答题'
        elif len(answer) > 0:
            return '填空题'
        
        return '未知'

